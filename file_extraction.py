import json
import os
import re
import zipfile
from dataclasses import dataclass
from io import BytesIO
from typing import Dict, Set
from xml.etree import ElementTree


SUPPORTED_FILE_TYPES: Dict[str, Dict[str, Set[str]]] = {
    ".txt": {"mime_types": {"text/plain", "application/octet-stream"}},
    ".md": {"mime_types": {"text/markdown", "text/plain", "application/octet-stream"}},
    ".csv": {"mime_types": {"text/csv", "application/csv", "text/plain", "application/octet-stream"}},
    ".json": {"mime_types": {"application/json", "text/json", "text/plain", "application/octet-stream"}},
    ".pdf": {"mime_types": {"application/pdf", "application/octet-stream"}},
    ".docx": {
        "mime_types": {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/octet-stream",
        }
    },
}

_GENERIC_BINARY_MIME_TYPES = {"application/octet-stream", "binary/octet-stream"}
_PDF_TJ_RE = re.compile(r"\(([^()]*)\)\s*Tj")
_PDF_TJ_ARRAY_RE = re.compile(r"\[(.*?)\]\s*TJ", flags=re.S)
_PDF_TJ_ARRAY_ITEM_RE = re.compile(r"\(([^()]*)\)")


class FileValidationError(ValueError):
    pass


class TextExtractionError(ValueError):
    pass


@dataclass(frozen=True)
class FileInfo:
    filename: str
    extension: str
    content_type: str


def _normalize_content_type(content_type: str) -> str:
    raw = str(content_type or "").strip().lower()
    if not raw:
        return "application/octet-stream"
    return raw.split(";", 1)[0].strip() or "application/octet-stream"


def validate_file_info(filename: str, content_type: str) -> FileInfo:
    safe_filename = os.path.basename(str(filename or "").strip())
    if not safe_filename:
        raise FileValidationError("filename is required")
    extension = os.path.splitext(safe_filename)[1].lower()
    if extension not in SUPPORTED_FILE_TYPES:
        supported = ", ".join(sorted(SUPPORTED_FILE_TYPES.keys()))
        raise FileValidationError(f"Unsupported file extension '{extension or '(none)'}'. Supported: {supported}")

    normalized_content_type = _normalize_content_type(content_type)
    allowed_mime_types = set(SUPPORTED_FILE_TYPES[extension]["mime_types"])
    if (
        normalized_content_type not in allowed_mime_types
        and normalized_content_type not in _GENERIC_BINARY_MIME_TYPES
    ):
        allowed_text = ", ".join(sorted(allowed_mime_types))
        raise FileValidationError(
            f"Unsupported MIME type '{normalized_content_type}' for '{extension}'. Allowed: {allowed_text}"
        )

    return FileInfo(
        filename=safe_filename,
        extension=extension,
        content_type=normalized_content_type,
    )


def _decode_text_bytes(raw: bytes) -> str:
    if not raw:
        return ""
    tried = []
    for encoding in ("utf-8-sig", "utf-16", "utf-16-le", "utf-16-be"):
        try:
            return raw.decode(encoding)
        except Exception:
            tried.append(encoding)
    if b"\x00" in raw:
        raise TextExtractionError("Binary content detected in text file.")
    try:
        return raw.decode("utf-8")
    except Exception as exc:
        raise TextExtractionError(f"Failed to decode text file (tried {', '.join(tried)} and utf-8).") from exc


def _extract_json_text(raw: bytes) -> str:
    decoded = _decode_text_bytes(raw)
    try:
        parsed = json.loads(decoded)
    except Exception as exc:
        raise TextExtractionError("Failed to decode JSON content.") from exc
    if isinstance(parsed, str):
        return parsed
    return json.dumps(parsed, ensure_ascii=False, indent=2)


def _extract_pdf_text_with_library(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(raw))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages).strip()


def _decode_pdf_token(token: str) -> str:
    return (
        token.replace(r"\(", "(")
        .replace(r"\)", ")")
        .replace(r"\n", "\n")
        .replace(r"\r", "\r")
        .replace(r"\t", "\t")
        .replace(r"\\", "\\")
    )


def _extract_pdf_text_fallback(raw: bytes) -> str:
    text = raw.decode("latin-1", errors="ignore")
    chunks = []
    for match in _PDF_TJ_RE.finditer(text):
        decoded = _decode_pdf_token(match.group(1))
        if decoded.strip():
            chunks.append(decoded)
    for array_match in _PDF_TJ_ARRAY_RE.finditer(text):
        fragment = array_match.group(1)
        items = [_decode_pdf_token(m.group(1)) for m in _PDF_TJ_ARRAY_ITEM_RE.finditer(fragment)]
        joined = "".join(items).strip()
        if joined:
            chunks.append(joined)
    return "\n".join(chunks).strip()


def _extract_pdf_text(raw: bytes) -> str:
    if not raw.startswith(b"%PDF"):
        raise TextExtractionError("Failed to extract text from PDF: invalid PDF file.")
    text = ""
    try:
        text = _extract_pdf_text_with_library(raw)
    except Exception:
        text = _extract_pdf_text_fallback(raw)
    if not text.strip():
        raise TextExtractionError("No extractable text found in PDF. OCR not yet enabled for image-only PDFs.")
    return text


def _extract_docx_text_with_library(raw: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(raw))
    lines = []
    for paragraph in document.paragraphs:
        value = str(paragraph.text or "").strip()
        if value:
            lines.append(value)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(str(cell.text or "").strip() for cell in row.cells if str(cell.text or "").strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines).strip()


def _extract_docx_text_fallback(raw: bytes) -> str:
    try:
        with zipfile.ZipFile(BytesIO(raw)) as archive:
            xml_payload = archive.read("word/document.xml")
    except Exception as exc:
        raise TextExtractionError("Failed to open DOCX file.") from exc
    try:
        root = ElementTree.fromstring(xml_payload)
    except Exception as exc:
        raise TextExtractionError("Failed to parse DOCX XML content.") from exc
    lines = [str(node.text).strip() for node in root.iter() if node.tag.endswith("}t") and str(node.text or "").strip()]
    return "\n".join(lines).strip()


def _extract_docx_text(raw: bytes) -> str:
    try:
        text = _extract_docx_text_with_library(raw)
    except Exception:
        text = _extract_docx_text_fallback(raw)
    if not text.strip():
        raise TextExtractionError("No extractable text found in DOCX file.")
    return text


def extract_text_from_bytes(raw: bytes, file_info: FileInfo) -> str:
    extension = file_info.extension
    if extension in {".txt", ".md", ".csv"}:
        return _decode_text_bytes(raw)
    if extension == ".json":
        return _extract_json_text(raw)
    if extension == ".pdf":
        return _extract_pdf_text(raw)
    if extension == ".docx":
        return _extract_docx_text(raw)
    raise TextExtractionError(f"Unsupported file extension '{extension}'.")
